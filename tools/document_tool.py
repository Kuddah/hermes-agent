#!/usr/bin/env python3
"""
Document Tool Module — HTML → DOCX / PDF document production.

Provides five actions:
  document_create   — Write HTML or Markdown source for a document
  document_convert  — Export to PDF (WeasyPrint), DOCX, Markdown, or TXT
  document_view     — Read document content and metadata
  document_list     — List documents in a project
  document_restore  — Restore a previous auto-versioned snapshot

Project layout (under HERMES_DELIVERABLES_DIR, default ./mnt):
  mnt/<project>/documents/<doc>.source.html   — canonical HTML source
  mnt/<project>/documents/<doc>.pdf / .docx   — exports
  mnt/<project>/documents/.<doc>.bak.<ts>.html — auto-snapshots

All file writes are security-checked via agent.file_safety.is_write_denied().
Paths are sandboxed under HERMES_DELIVERABLES_DIR (no .. traversal allowed).

Unicode normalization for PDF: smart quotes, em/en dashes, ellipsis and their
truncated low-byte forms are mapped to ASCII equivalents to prevent font-glyph
failures in WeasyPrint.
"""

from __future__ import annotations

import json
import logging
import os
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from tools.html_validation import find_unsupported_html, build_unsupported_error

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# HERMES_DELIVERABLES_DIR — configurable sandbox root
# ---------------------------------------------------------------------------

def _get_deliverables_dir() -> Path:
    """Return the root directory for all deliverable projects."""
    raw = os.environ.get("HERMES_DELIVERABLES_DIR", "")
    if raw:
        return Path(raw).expanduser().resolve()
    # Docker: /app/mnt; local: ./mnt next to run_agent.py
    if Path("/.dockerenv").is_file():
        return Path("/app/mnt")
    return Path.cwd() / "mnt"


def _get_project_dir(project_name: str) -> Path:
    return _get_deliverables_dir() / _safe_name(project_name) / "documents"


def _safe_name(name: str) -> str:
    """Strip path separators from project/doc names to prevent traversal."""
    return re.sub(r"[/\\]", "_", name.strip())


# ---------------------------------------------------------------------------
# Security helpers
# ---------------------------------------------------------------------------

def _check_write(path: Path) -> Optional[str]:
    """Return an error string if the path is write-denied, else None."""
    from agent.file_safety import is_write_denied
    if is_write_denied(str(path)):
        return f"Write denied by security policy: {path}"
    # Ensure path stays inside deliverables root
    root = _get_deliverables_dir().resolve()
    try:
        path.resolve().relative_to(root)
    except ValueError:
        return f"Path escapes deliverables sandbox: {path}"
    # Enforce extension allowlist
    ext_err = _validate_extension(path)
    if ext_err:
        return ext_err
    return None


_ALLOWED_EXTENSIONS = {".html", ".docx", ".pdf", ".md", ".txt", ".png", ".jpg"}


def _validate_extension(path: Path) -> Optional[str]:
    if path.suffix.lower() not in _ALLOWED_EXTENSIONS:
        return f"Extension '{path.suffix}' not allowed. Permitted: {', '.join(sorted(_ALLOWED_EXTENSIONS))}"
    return None


# ---------------------------------------------------------------------------
# Unicode normalization (prevents PDF font-glyph failures)
# ---------------------------------------------------------------------------

_UNICODE_TO_ASCII = str.maketrans({
    "\u2018": "'",  "\u2019": "'",   # ' '
    "\u201c": '"',  "\u201d": '"',   # " "
    "\u2013": "-",  "\u2014": "--",  # – —
    "\u2026": "...","\u00a0": " ",   # … non-breaking space
    "\x19": "'", "\x11": "-", "\x18": "'", "\x1c": '"', "\x1d": '"', "\x14": "--",
})


def _normalize_unicode(text: str) -> str:
    return text.translate(_UNICODE_TO_ASCII)


# ---------------------------------------------------------------------------
# Snapshot helpers
# ---------------------------------------------------------------------------

def _snapshot(source_path: Path) -> None:
    """Copy source_path to a timestamped backup beside it (silent on error)."""
    try:
        ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
        backup = source_path.parent / f".{source_path.stem}.bak.{ts}{source_path.suffix}"
        shutil.copy2(source_path, backup)
    except Exception:
        pass


def _prune_old_snapshots(project_dir: Path, doc_name: str, keep_days: int = 7) -> None:
    """Remove auto-snapshots older than keep_days."""
    try:
        cutoff = datetime.now(timezone.utc).timestamp() - keep_days * 86400
        for f in project_dir.glob(f".{doc_name}.bak.*"):
            if f.stat().st_mtime < cutoff:
                f.unlink(missing_ok=True)
    except Exception:
        pass


def _list_snapshots(project_dir: Path, doc_name: str) -> list[Path]:
    return sorted(project_dir.glob(f".{doc_name}.bak.*"), key=lambda p: p.stat().st_mtime)


# ---------------------------------------------------------------------------
# Versioning helper (like OpenSwarm's next_docx_version)
# ---------------------------------------------------------------------------

def _next_version(desired: Path) -> Path:
    if not desired.exists():
        return desired
    base = re.sub(r"_v\d+$", "", desired.stem)
    n = 2
    while True:
        candidate = desired.parent / f"{base}_v{n}{desired.suffix}"
        if not candidate.exists():
            return candidate
        n += 1


# ---------------------------------------------------------------------------
# Dependency availability checks
# ---------------------------------------------------------------------------

def check_document_requirements() -> bool:
    """Return True if at least the write_file path is available (always True)."""
    return True


def _weasyprint_available() -> bool:
    try:
        import weasyprint  # noqa: F401
        return True
    except ImportError:
        return False


def _docx_available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


# ---------------------------------------------------------------------------
# Core actions
# ---------------------------------------------------------------------------

def _create(project: str, doc_name: str, content: str, content_type: str,
            overwrite: bool) -> str:
    doc_name = _safe_name(doc_name.replace(".html", "").replace(".docx", "").replace(".md", ""))
    project_dir = _get_project_dir(project)

    if content_type == "markdown":
        md_path = project_dir / f"{doc_name}.md"
        if md_path.exists() and not overwrite:
            return json.dumps({"error": f"Document '{doc_name}.md' already exists. Use overwrite=true to replace."})
        err = _check_write(md_path)
        if err:
            return json.dumps({"error": err})
        project_dir.mkdir(parents=True, exist_ok=True)
        if md_path.exists():
            _snapshot(md_path)
        md_path.write_text(content, encoding="utf-8")
        return json.dumps({"ok": True, "path": str(md_path), "format": "markdown"})

    # HTML path
    source_path = project_dir / f"{doc_name}.source.html"
    if source_path.exists() and not overwrite:
        return json.dumps({"error": f"Document '{doc_name}' already exists. Use overwrite=true to replace."})
    err = _check_write(source_path)
    if err:
        return json.dumps({"error": err})
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "assets").mkdir(exist_ok=True)
    if source_path.exists():
        _snapshot(source_path)

    # Validate HTML for DOCX-export compatibility
    html_warnings = find_unsupported_html(content)

    source_path.write_text(content, encoding="utf-8")
    _prune_old_snapshots(project_dir, doc_name)
    result = {
        "ok": True,
        "path": str(source_path),
        "format": "html",
        "size_bytes": source_path.stat().st_size,
        "hint": f"Run document_convert with project='{project}' name='{doc_name}' format='pdf' to export.",
    }
    if html_warnings:
        result["docx_warnings"] = build_unsupported_error(html_warnings)
    return json.dumps(result)


def _convert(project: str, doc_name: str, output_format: str, overwrite: bool) -> str:
    doc_name = _safe_name(doc_name.replace(".html", "").replace(".docx", "").replace(".md", ""))
    project_dir = _get_project_dir(project)
    source_path = project_dir / f"{doc_name}.source.html"

    if not project_dir.exists():
        return json.dumps({"error": f"Project '{project}' not found."})
    if not source_path.exists():
        return json.dumps({"error": f"Document '{doc_name}' source not found in project '{project}'."})

    html_content = _normalize_unicode(source_path.read_text(encoding="utf-8"))

    ext_map = {"pdf": ".pdf", "docx": ".docx", "markdown": ".md", "txt": ".txt"}
    if output_format not in ext_map:
        return json.dumps({"error": f"Unknown format '{output_format}'. Choose from: {', '.join(ext_map)}."})

    output_path = _next_version(project_dir / f"{doc_name}{ext_map[output_format]}") \
        if output_format == "docx" \
        else project_dir / f"{doc_name}{ext_map[output_format]}"

    if output_path.exists() and not overwrite and output_format != "docx":
        return json.dumps({"error": f"Output '{output_path.name}' exists. Use overwrite=true."})

    err = _check_write(output_path)
    if err:
        return json.dumps({"error": err})

    try:
        if output_format == "pdf":
            if not _weasyprint_available():
                return json.dumps({"error": "WeasyPrint not installed. Run: pip install 'hermes-agent[deliverables]'"})
            from weasyprint import HTML as WP_HTML
            WP_HTML(string=html_content, base_url=str(project_dir)).write_pdf(str(output_path))

        elif output_format == "docx":
            if not _docx_available():
                return json.dumps({"error": "python-docx not installed. Run: pip install 'hermes-agent[deliverables]'"})
            # Warn about unsupported CSS features before DOCX conversion
            docx_issues = find_unsupported_html(html_content)
            _html_to_docx(html_content, output_path, project_dir)
            if docx_issues:
                return json.dumps({
                    "ok": True,
                    "output": str(output_path),
                    "size_bytes": output_path.stat().st_size,
                    "warnings": build_unsupported_error(docx_issues),
                })
                # Fall through to normal return below if no issues

        elif output_format == "markdown":
            try:
                import html2text
                h = html2text.HTML2Text()
                h.ignore_links = False
                output_path.write_text(h.handle(html_content), encoding="utf-8")
            except ImportError:
                # Fallback: strip tags
                import re as _re
                output_path.write_text(_re.sub(r"<[^>]+>", "", html_content), encoding="utf-8")

        elif output_format == "txt":
            try:
                from bs4 import BeautifulSoup
                output_path.write_text(BeautifulSoup(html_content, "html.parser").get_text(), encoding="utf-8")
            except ImportError:
                import re as _re
                output_path.write_text(_re.sub(r"<[^>]+>", "", html_content), encoding="utf-8")

        return json.dumps({
            "ok": True,
            "output": str(output_path),
            "size_bytes": output_path.stat().st_size,
        })
    except Exception as exc:
        logger.error("document_convert failed: %s", exc, exc_info=True)
        return json.dumps({"error": str(exc)})


def _html_to_docx(html: str, output_path: Path, base_dir: Path) -> None:
    """HTML→DOCX via python-docx + beautifulsoup with rich formatting.

    Preserves bold, italic, code spans, links (as text), images (as
    placeholders), and tables with headers — significantly better
    fidelity than the naive text-only extraction.
    """
    from bs4 import BeautifulSoup, Tag, NavigableString
    import docx
    from docx.shared import Pt, Inches

    doc = docx.Document()
    soup = BeautifulSoup(html, "html.parser")
    body = soup.find("body") or soup

    def _add_inline_runs(paragraph, element):
        """Recursively add inline runs preserving bold/italic/code formatting."""
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip() or text == " ":
                    paragraph.add_run(text)
            elif isinstance(child, Tag):
                if child.name in ("br",):
                    paragraph.add_run("\n")
                elif child.name in ("b", "strong"):
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name in ("i", "em"):
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                elif child.name in ("code",):
                    run = paragraph.add_run(child.get_text())
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                elif child.name == "a":
                    link_text = child.get_text()
                    href = child.get("href", "")
                    run = paragraph.add_run(link_text)
                    run.underline = True
                    if href and href != "#":
                        run = paragraph.add_run(f" ({href})")
                        run.font.size = Pt(8)
                elif child.name == "img":
                    alt = child.get("alt", "[image]")
                    run = paragraph.add_run(f"[{alt}]")
                    run.italic = True
                else:
                    _add_inline_runs(paragraph, child)

    def _process_table(table_tag):
        """Convert an HTML <table> to a DOCX table."""
        rows = table_tag.find_all("tr")
        if not rows:
            return
        # Determine column count from first row
        first_cells = rows[0].find_all(["th", "td"])
        ncols = len(first_cells)
        if ncols == 0:
            return
        tbl = doc.add_table(rows=0, cols=ncols)
        tbl.style = "Table Grid"
        for row_tag in rows:
            cells = row_tag.find_all(["th", "td"])
            row = tbl.add_row()
            for i, cell_tag in enumerate(cells):
                if i < ncols:
                    cell_text = cell_tag.get_text(separator=" ").strip()
                    row.cells[i].text = cell_text
                    if cell_tag.name == "th":
                        for run in row.cells[i].paragraphs[0].runs:
                            run.bold = True

    # Walk top-level block elements
    block_tags = {"h1", "h2", "h3", "h4", "h5", "h6", "p", "li",
                  "blockquote", "pre", "table", "ul", "ol", "div", "section"}

    def _walk_blocks(parent):
        for child in parent.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    p = doc.add_paragraph()
                    p.add_run(text)
                continue
            if not isinstance(child, Tag):
                continue

            name = child.name

            if name == "table":
                _process_table(child)
            elif name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                level = int(name[1])
                doc.add_heading(child.get_text(separator=" ").strip(), level=min(level, 4))
            elif name == "p":
                p = doc.add_paragraph()
                _add_inline_runs(p, child)
            elif name in ("ul", "ol"):
                for li in child.find_all("li", recursive=False):
                    p = doc.add_paragraph(style="List Bullet" if name == "ul" else "List Number")
                    _add_inline_runs(p, li)
            elif name == "li":
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, child)
            elif name == "blockquote":
                text = child.get_text(separator=" ").strip()
                p = doc.add_paragraph(text)
                p.paragraph_format.left_indent = Inches(0.5)
                for run in p.runs:
                    run.italic = True
            elif name == "pre":
                code_text = child.get_text()
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            elif name in ("div", "section", "main", "article", "header", "footer"):
                _walk_blocks(child)
            else:
                # Unknown block — try to extract text
                text = child.get_text(separator=" ").strip()
                if text:
                    p = doc.add_paragraph()
                    _add_inline_runs(p, child)

    _walk_blocks(body)
    doc.save(str(output_path))


def _view(project: str, doc_name: str) -> str:
    doc_name = _safe_name(doc_name.replace(".html", "").replace(".docx", "").replace(".md", ""))
    project_dir = _get_project_dir(project)
    source_path = project_dir / f"{doc_name}.source.html"
    md_path = project_dir / f"{doc_name}.md"

    if source_path.exists():
        content = source_path.read_text(encoding="utf-8")
        exports = [p.name for p in project_dir.glob(f"{doc_name}.*") if p.suffix != ".html"]
        return json.dumps({
            "project": project, "name": doc_name, "format": "html",
            "size_bytes": source_path.stat().st_size,
            "content": content[:8000] + ("... [truncated]" if len(content) > 8000 else ""),
            "exports": exports,
        })
    elif md_path.exists():
        content = md_path.read_text(encoding="utf-8")
        return json.dumps({
            "project": project, "name": doc_name, "format": "markdown",
            "size_bytes": md_path.stat().st_size,
            "content": content[:8000] + ("... [truncated]" if len(content) > 8000 else ""),
        })
    else:
        return json.dumps({"error": f"Document '{doc_name}' not found in project '{project}'."})


def _list_docs(project: str) -> str:
    project_dir = _get_project_dir(project)
    if not project_dir.exists():
        return json.dumps({"project": project, "documents": []})
    docs = []
    for src in sorted(project_dir.glob("*.source.html")):
        name = src.name.replace(".source.html", "")
        exports = [p.name for p in project_dir.glob(f"{name}.*") if p.suffix != ".html"]
        docs.append({"name": name, "exports": exports, "size_bytes": src.stat().st_size})
    for md in sorted(project_dir.glob("*.md")):
        if not md.name.startswith("."):
            docs.append({"name": md.stem, "format": "markdown", "size_bytes": md.stat().st_size})
    return json.dumps({"project": project, "documents": docs})


def _restore(project: str, doc_name: str, snapshot_index: int) -> str:
    doc_name = _safe_name(doc_name.replace(".html", "").replace(".docx", "").replace(".md", ""))
    project_dir = _get_project_dir(project)
    snapshots = _list_snapshots(project_dir, doc_name)
    if not snapshots:
        return json.dumps({"error": f"No snapshots found for '{doc_name}'."})
    idx = snapshot_index - 1
    if idx < 0 or idx >= len(snapshots):
        available = [f"{i+1}: {p.name}" for i, p in enumerate(snapshots)]
        return json.dumps({"error": "snapshot_index out of range.", "available": available})
    snap = snapshots[idx]
    target = project_dir / f"{doc_name}.source.html"
    err = _check_write(target)
    if err:
        return json.dumps({"error": err})
    _snapshot(target)  # back up current before overwriting
    shutil.copy2(snap, target)
    return json.dumps({"ok": True, "restored_from": snap.name, "target": str(target)})


# ---------------------------------------------------------------------------
# Unified dispatcher
# ---------------------------------------------------------------------------

def document_tool(action: str, **kwargs) -> str:
    """Dispatch document tool actions."""
    if action == "create":
        return _create(
            project=kwargs.get("project", ""),
            doc_name=kwargs.get("name", ""),
            content=kwargs.get("content", ""),
            content_type=kwargs.get("content_type", "html"),
            overwrite=bool(kwargs.get("overwrite", False)),
        )
    elif action == "convert":
        return _convert(
            project=kwargs.get("project", ""),
            doc_name=kwargs.get("name", ""),
            output_format=kwargs.get("format", "pdf"),
            overwrite=bool(kwargs.get("overwrite", True)),
        )
    elif action == "view":
        return _view(kwargs.get("project", ""), kwargs.get("name", ""))
    elif action == "list":
        return _list_docs(kwargs.get("project", ""))
    elif action == "restore":
        return _restore(
            project=kwargs.get("project", ""),
            doc_name=kwargs.get("name", ""),
            snapshot_index=int(kwargs.get("snapshot_index", 1)),
        )
    else:
        return json.dumps({"error": f"Unknown action '{action}'. Use: create, convert, view, list, restore"})


# =============================================================================
# OpenAI Function-Calling Schema
# =============================================================================

DOCUMENT_SCHEMA = {
    "name": "document_tool",
    "description": (
        "Create, export, view, and manage rich documents (HTML→PDF/DOCX/Markdown).\n\n"
        "Actions:\n"
        "- **create**: Write an HTML or Markdown document. HTML is the canonical source; "
        "always create HTML first, then convert.\n"
        "- **convert**: Export the HTML source to PDF (WeasyPrint), DOCX (python-docx), "
        "Markdown, or TXT.\n"
        "- **view**: Read document content and list available exports.\n"
        "- **list**: List all documents in a project.\n"
        "- **restore**: Revert to an auto-saved snapshot.\n\n"
        "All outputs go to mnt/<project>/documents/. Use lowercase_underscores for names.\n"
        "HTML-first workflow: create HTML → iterate → convert to final format."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "enum": ["create", "convert", "view", "list", "restore"],
                "description": "Operation to perform.",
            },
            "project": {
                "type": "string",
                "description": "Project folder name (e.g. 'quarterly_report'). Use lowercase_underscores.",
            },
            "name": {
                "type": "string",
                "description": "Document name without extension (e.g. 'executive_summary').",
            },
            "content": {
                "type": "string",
                "description": "Full HTML or Markdown content for the 'create' action.",
            },
            "content_type": {
                "type": "string",
                "enum": ["html", "markdown"],
                "default": "html",
                "description": "Content format: 'html' (default, enables PDF/DOCX export) or 'markdown'.",
            },
            "format": {
                "type": "string",
                "enum": ["pdf", "docx", "markdown", "txt"],
                "description": "Export format for 'convert' action.",
            },
            "overwrite": {
                "type": "boolean",
                "default": False,
                "description": "For 'create': overwrite existing source. For 'convert': overwrite existing export.",
            },
            "snapshot_index": {
                "type": "integer",
                "description": "1-based snapshot index for 'restore'. Call with action='list' first to see snapshots.",
            },
        },
        "required": ["action"],
    },
}


# --- Registry ---
from tools.registry import registry, tool_error  # noqa: E402

registry.register(
    name="document_tool",
    toolset="document",
    schema=DOCUMENT_SCHEMA,
    handler=lambda args, **kw: document_tool(
        action=args.get("action", ""),
        project=args.get("project", ""),
        name=args.get("name", ""),
        content=args.get("content", ""),
        content_type=args.get("content_type", "html"),
        format=args.get("format", "pdf"),
        overwrite=args.get("overwrite", False),
        snapshot_index=args.get("snapshot_index", 1),
    ),
    check_fn=check_document_requirements,
    emoji="📄",
)
